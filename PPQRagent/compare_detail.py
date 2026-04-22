"""각 표별 상세 비교 및 원인 분석"""
import sys
from docx import Document
from docx.oxml.ns import qn
import openpyxl

out = sys.stdout
enc = "utf-8"
W = lambda s: out.buffer.write((str(s) + "\n").encode(enc, "replace"))

human = Document("PQR-AMP-26-002 이달비정80밀리그램.docx")
auto  = Document("PPQR_완성본.docx")

# ─────────────────────────────────────────────────────────
W("\n" + "="*60)
W("[분석1] T11 - 실제 셀 구조 확인 (병합 여부)")
W("="*60)
tbl = human.tables[10]
W(f"사람버전 T11: {len(tbl.rows)}행 x {len(tbl.columns)}열")
# 병합 셀 감지
from docx.oxml.ns import qn
def get_span(cell):
    tc = cell._tc
    vmerge = tc.find(qn("w:tcPr") + "/" + qn("w:vMerge"))
    gridSpan = tc.find(qn("w:tcPr") + "/" + qn("w:gridSpan"))
    vspan = tc.find(".//" + qn("w:vMerge"))
    gspan = tc.find(".//" + qn("w:gridSpan"))
    return {
        "vmerge": vspan.get(qn("w:val"), "continue") if vspan is not None else None,
        "gridspan": gspan.get(qn("w:val"), "1") if gspan is not None else None,
    }

W("\n사람버전 T11 각 행 셀 내용 (실제 XML 기반):")
for ri, row in enumerate(tbl.rows[:6]):
    row_data = []
    for ci, cell in enumerate(row.cells):
        sp = get_span(cell)
        text = cell.text.strip().replace("\n", "/")[:25]
        row_data.append(f"[{ci}]{text!r}(v={sp['vmerge']},g={sp['gridspan']})")
    W(f"  row{ri}: " + " | ".join(row_data[:5]))

# ─────────────────────────────────────────────────────────
W("\n" + "="*60)
W("[분석2] T16 Close date - Excel 원본 확인")
W("="*60)
wb = openpyxl.load_workbook("sample/5) 변경관리 목록_2025년.xlsx", data_only=True)
ws = wb["Change Control List"]
TARGET = ["30234","30497","3023401","이달비정 80","이달비정80","Edarbi 80"]
W("CC 파일에서 해당 행 날짜 컬럼 (col 11=Pre-Approval, col 12=Completion):")
for r in range(5, 660):
    if not ws.cell(r,2).value: continue
    affected = str(ws.cell(r,5).value or "")
    if not any(t in affected for t in TARGET): continue
    cc_no   = str(ws.cell(r,2).value or "")
    col11   = ws.cell(r,11).value
    col12   = ws.cell(r,12).value
    col13   = ws.cell(r,13).value  # Remark?
    W(f"  row{r} | {cc_no} | col11={col11} | col12={col12} | col13={col13}")

# ─────────────────────────────────────────────────────────
W("\n" + "="*60)
W("[분석3] T18 Close date - Excel 원본 확인")
W("="*60)
wb2 = openpyxl.load_workbook("sample/6) List of CAPA (2025)_QA1.xlsx", data_only=True)
ws2 = wb2["Chemical"]
TARGET2 = ["30234","30497","3023401","이달비정 80","이달비정80","Edarbi 80"]
W("CAPA 파일 컬럼 (1=CAPA No, 6=Status, 7=Opened, 8=Planned, 9=Actual?):")
# 헤더 확인
hdrs = [str(ws2.cell(2, c).value or "")[:20] for c in range(1, 12)]
W(f"  헤더(row2): {hdrs}")
for r in range(3, ws2.max_row+1):
    capa_no = str(ws2.cell(r,1).value or "").strip()
    if not capa_no or "CAPA" not in capa_no: continue
    product = str(ws2.cell(r,5).value or "")
    if not any(t in f"{capa_no} {product}" for t in TARGET2): continue
    # 모든 날짜 관련 컬럼 출력
    vals = [str(ws2.cell(r,c).value or "")[:15] for c in range(1,12)]
    W(f"  row{r}: {vals}")

# ─────────────────────────────────────────────────────────
W("\n" + "="*60)
W("[분석4] T11 Packaging Unit - Excel 원본")
W("="*60)
wb3 = openpyxl.load_workbook("sample/2) 자사 제조배치.xlsx", data_only=True, read_only=True)
ws3 = wb3.active
CODES = ["3023401","3023402","3023403","3023404"]
W("배치 파일 col 3=포장단위 실제값:")
for r, rv in enumerate(ws3.iter_rows(min_row=4, max_row=30, values_only=True), start=4):
    if not rv[0]: continue
    code = str(rv[2] or "")
    if not any(c in code for c in CODES): continue
    W(f"  row{r} | code={code} | col3(포장단위)={rv[3]!r} | col4(제조번호)={rv[4]!r}")
wb3.close()

# ─────────────────────────────────────────────────────────
W("\n" + "="*60)
W("[분석5] T19 불만 - 사람버전 포맷 분석")
W("="*60)
tbl19h = human.tables[18]
tbl19a = auto.tables[18]
W("사람버전 T19 데이터 셀 (col2=Contents, col3=Date):")
for ri, row in enumerate(tbl19h.rows[1:4]):
    cells = list(row.cells)
    # 중복 제거
    seen = set()
    unique = []
    for c in cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            unique.append(c)
    W(f"  row{ri}: {[c.text.strip()[:50] for c in unique]}")
W("\n자동생성 T19:")
for ri, row in enumerate(tbl19a.rows[1:4]):
    seen = set(); unique = []
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc)); unique.append(c)
    W(f"  row{ri}: {[c.text.strip()[:50] for c in unique]}")

# ─────────────────────────────────────────────────────────
W("\n" + "="*60)
W("[분석6] T19 - Excel 원본 컬럼 구조")
W("="*60)
wb4 = openpyxl.load_workbook("sample/7) 2025년도 소비자불만 리스트.xlsx", data_only=True)
ws4 = wb4.active
hdrs4 = [str(ws4.cell(1,c).value or "")[:15] for c in range(1, 20)]
W(f"헤더: {hdrs4}")
TARGET4 = ["이달비정 80","이달비정80","Edarbi 80mg","Edarbi Tab. 80","30234","30497"]
for r in range(2, ws4.max_row+1):
    product = str(ws4.cell(r,5).value or "")
    if not any(t in product for t in TARGET4): continue
    # 전체 컬럼 출력
    vals = [str(ws4.cell(r,c).value or "")[:20] for c in range(1, 20)]
    W(f"  row{r}: {vals}")
