"""
PPQR_빈양식.docx 범용화:
- 표 밖 단락 텍스트 → 빨간색
- 제품명 → "xx xxmg"
- 문서번호 → "xx-xx"
- 헤더/푸터도 동일 처리
"""

from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import re

SRC = "d:/workAnti/PPQRagent/PPQR_빈양식.docx"
DST = "d:/workAnti/PPQRagent/PPQR_빈양식.docx"

RED = RGBColor(0xFF, 0x00, 0x00)

# 제품명 패턴 — 단락 전체 텍스트에 적용
PRODUCT_PATTERNS = [
    (r"이달비정\s*80\s*밀리그램\s*[\(\（][^\)\）]*[\)\）]", "xx xxmg"),
    (r"이달비정\s*80\s*밀리그램", "xx xxmg"),
    (r"Edarbi\s+Tab\.\s*80\s*mg", "xx xxmg"),
    (r"Edarbi\s+80\s*mg", "xx xxmg"),
    (r"Edarbi", "xx xxmg"),
    (r"Azilsartan\s+Medoxomil\s+Potassium", "xx xxmg"),
    (r"Azilsartan\s+Medoxomil", "xx xxmg"),
    (r"Azilsartan", "xx xxmg"),
    (r"아질사르탄메독소밀칼륨", "xx xxmg"),
    (r"아질사르탄메독소밀", "xx xxmg"),
    (r"PQR-AMP-\d{2}-\d{3}", "xx-xx"),
]


def apply_replacements(text: str) -> str:
    for pat, repl in PRODUCT_PATTERNS:
        text = re.sub(pat, repl, text, flags=re.IGNORECASE)
    return text


def process_paragraph(para: Paragraph, make_red: bool = True):
    """
    단락 전체 텍스트를 합쳐서 치환한 뒤 첫 번째 run에 기록.
    나머지 run은 텍스트를 비워 중복 방지.
    make_red=True 면 모든 run 빨간색으로.
    """
    runs = para.runs
    if not runs:
        return

    # 전체 텍스트 합산 후 치환
    full_text = para.text
    replaced = apply_replacements(full_text)

    if replaced != full_text or make_red:
        # 첫 run에 전체 대체 텍스트 기록, 나머지 run 비우기
        runs[0].text = replaced
        for r in runs[1:]:
            r.text = ""
        # 색상 적용
        if make_red:
            runs[0].font.color.rgb = RED
        elif replaced != full_text:
            # 표 안 run은 색 변경 없이 텍스트만 치환
            pass


def iter_body_top_children(doc):
    """body 직계 자식 순회 — 표 밖 단락과 표 구분"""
    body = doc.element.body
    for child in body:
        yield child


doc = Document(SRC)

# ── 본문 처리 ─────────────────────────────────────────────────────
for child in iter_body_top_children(doc):
    if child.tag == qn("w:p"):
        para = Paragraph(child, doc)
        process_paragraph(para, make_red=True)
    elif child.tag == qn("w:tbl"):
        # 표 안: 텍스트 치환만 (빨간색 X)
        for tc in child.iter(qn("w:tc")):
            for p_elem in tc.iter(qn("w:p")):
                para = Paragraph(p_elem, doc)
                process_paragraph(para, make_red=False)

# ── 헤더/푸터 처리 ───────────────────────────────────────────────
for section in doc.sections:
    for hdr in [
        section.header, section.footer,
        section.even_page_header, section.even_page_footer,
        section.first_page_header, section.first_page_footer,
    ]:
        if hdr is None:
            continue
        try:
            for para in hdr.paragraphs:
                process_paragraph(para, make_red=True)
            for tbl in hdr.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            process_paragraph(para, make_red=False)
        except Exception:
            pass

doc.save(DST)
print("저장 완료:", DST)
