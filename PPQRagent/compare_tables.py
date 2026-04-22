"""두 Word 파일의 핵심 표를 비교 분석"""
import sys
from docx import Document

def extract_table(doc, tbl_idx):
    tbl = doc.tables[tbl_idx]
    rows = []
    seen = set()
    for row in tbl.rows:
        cells = []
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen:
                cells.append("(merged)")
            else:
                seen.add(cid)
                cells.append(cell.text.strip().replace("\n", " / "))
        rows.append(cells)
    return rows

KEY_TABLES = {10: "T11 제조배치", 15: "T16 변경관리", 17: "T18 CAPA", 18: "T19 불만"}

human = Document("PQR-AMP-26-002 이달비정80밀리그램.docx")
auto  = Document("PPQR_완성본.docx")

out = sys.stdout
enc = "utf-8"

for tbl_idx, name in KEY_TABLES.items():
    out.buffer.write(f"\n{'='*70}\n[{name}]  Table[{tbl_idx}]\n{'='*70}\n".encode(enc))

    h_rows = extract_table(human, tbl_idx)
    a_rows = extract_table(auto,  tbl_idx)

    out.buffer.write(f"  사람: {len(h_rows)}행 (헤더 포함)  /  자동: {len(a_rows)}행\n\n".encode(enc))

    out.buffer.write("  [사람 작업 - 헤더]\n".encode(enc))
    if h_rows:
        out.buffer.write(("  " + " | ".join(h_rows[0][:8]) + "\n").encode(enc, "replace"))

    out.buffer.write("\n  [사람 작업 - 데이터 행]\n".encode(enc))
    for r in h_rows[1:8]:
        line = "  " + " | ".join(str(c)[:30] for c in r[:8])
        out.buffer.write((line + "\n").encode(enc, "replace"))

    out.buffer.write("\n  [자동생성 - 데이터 행]\n".encode(enc))
    for r in a_rows[1:8]:
        line = "  " + " | ".join(str(c)[:30] for c in r[:8])
        out.buffer.write((line + "\n").encode(enc, "replace"))

    # 컬럼별 내용 유무 비교
    out.buffer.write("\n  [빈 셀 현황]\n".encode(enc))
    for ri, (hrow, arow) in enumerate(zip(h_rows[1:], a_rows[1:]), start=1):
        diffs = []
        for ci, (hc, ac) in enumerate(zip(hrow, arow)):
            if hc and not ac:
                diffs.append(f"  col{ci}: 사람={hc[:20]!r} / 자동=빈값")
            elif hc != ac and ac:
                diffs.append(f"  col{ci}: 사람={hc[:20]!r} / 자동={ac[:20]!r}")
        if diffs:
            out.buffer.write(f"\n  row{ri}:\n".encode(enc))
            for d in diffs[:5]:
                out.buffer.write((d + "\n").encode(enc, "replace"))
