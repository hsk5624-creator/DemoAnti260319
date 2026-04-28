import os, re, json, zipfile, xml.etree.ElementTree as ET, copy
from pathlib import Path
from datetime import datetime
from contextlib import asynccontextmanager
import asyncio
import openpyxl
from fastapi import FastAPI, Query, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

@asynccontextmanager
async def lifespan(app: FastAPI):
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, _preload_cache)
    yield

app = FastAPI(lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

BASE    = Path(__file__).parent.parent
SAMPLE  = BASE / "sample"
UPLOADS = BASE / "uploads"
UPLOADS.mkdir(exist_ok=True)
PRODUCT_INFO_FILE = UPLOADS / "product_info.json"

SAMPLE_FILES = {
    "batch":     SAMPLE / "2) 자사 제조배치.xlsx",
    "wms":       SAMPLE / "2) 2025년도 WMS 제조지시 리스트_20260109_제조일자 확인.xlsx",
    "cc":        SAMPLE / "5) 변경관리 목록_2025년.xlsx",
    "capa":      SAMPLE / "6) List of CAPA (2025)_QA1.xlsx",
    "complaint": SAMPLE / "7) 2025년도 소비자불만 리스트.xlsx",
    "scar":      SAMPLE / "10) QS06012-A4(1.0) Tracking Database for Supplier Corrective Action_20260113 (최신).xlsx",
    "deviation": SAMPLE / "4) QS09001-A3 Deviation Log 2025.xlsx",
    "banpum":    SAMPLE / "7) 반품 내역.xlsb",
    "rmt":       SAMPLE / "10) 원자재 시험성적 관리.xlsb",
    "fpt":       SAMPLE / "2) 완제품 시험성적 관리.xlsb",
    "bdc_100k":  SAMPLE / "Batch Data Collection_300244_M04 80mg (Inist)_O.xlsx",
    "bdc_700k":  SAMPLE / "Batch Data Collection_300244_M04 80mg (Inist) (Batch size 700,000 T)_O.xlsx",
    "template":  BASE   / "제품명 C-PPQR-허가번호-년도(버전)_작성예시2025.docx",
}

UPLOAD_FILES = {
    "batch":     UPLOADS / "batch.xlsx",
    "wms":       UPLOADS / "wms.xlsx",
    "cc":        UPLOADS / "cc.xlsx",
    "capa":      UPLOADS / "capa.xlsx",
    "complaint": UPLOADS / "complaint.xlsx",
    "scar":      UPLOADS / "scar.xlsx",
    "deviation": UPLOADS / "deviation.xlsx",
    "banpum":    UPLOADS / "banpum.xlsb",
    "rmt":       UPLOADS / "rmt.xlsb",
    "fpt":       UPLOADS / "fpt.xlsb",
    "template":  UPLOADS / "template.docx",
}

def get_path(key: str) -> Path:
    up = UPLOAD_FILES.get(key)
    if up and up.exists():
        return up
    return SAMPLE_FILES[key]

# kept for 이달비정80 display fallback (falls back to raw code if product not in map)
MARKET_MAP = {"3023401": "KR", "3023402": "HK", "3023403": "TH", "3023404": "TH"}
PACKAGING_UNIT_MAP = {
    "3023401": "28T/\nAlu-Alu\n(14T x 2ea)",
    "3023402": "28T/\nAlu-Alu\n(14T x 2ea)",
    "3023403": "28T/\nAlu-Alu\n(14T x 2ea)",
    "3023404": "28T/\nAlu-Alu\n(14T x 2ea)",
}
BATCH_NORMAL_RE = re.compile(r"^(CJ|CL|CH)[A-Z]{2}\d{3}$")

COMPLAINT_TYPE_EN: dict[str, str] = {
    "내용량부족": "Insufficient content",
    "포장 불량(포켓 불량)": "Packaging defect (Damaged pocket)",
    "포장불량": "Packaging defect",
    "기타(포장 오염)": "Others (Packaging contamination – Mold)",
    "기타(포장오염)": "Others (Packaging contamination – Mold)",
    "기타": "Others",
    "이물": "Foreign matter",
    "변색": "Discoloration",
    "변취": "Odor change",
    "균열/파쇄": "Crack/Breakage",
    "외관불량": "Appearance defect",
    "인쇄불량": "Print defect",
}

def load_wb(key, read_only=False):
    return openpyxl.load_workbook(get_path(key), data_only=True, read_only=read_only)

# ── 제품 설정 ─────────────────────────────────────────────────────────
def get_product_config(product_name: str = "", product_code: str = "", api_name: str = "") -> dict:
    """Build product filter config from params or saved file."""
    if not product_name and not product_code:
        if PRODUCT_INFO_FILE.exists():
            try:
                info = json.loads(PRODUCT_INFO_FILE.read_text(encoding="utf-8"))
                product_name = info.get("product_name", "")
                product_code = info.get("product_code", "")
                api_name     = info.get("api_name", "")
            except Exception:
                pass

    name = product_name.strip()
    code = product_code.strip()
    api  = api_name.strip()

    # Generate name variants: original + no-space + space-before-trailing-number
    name_variants: list[str] = []
    if name:
        name_variants.append(name)
        no_space = name.replace(" ", "")
        if no_space != name:
            name_variants.append(no_space)
        m = re.match(r'^(.*\D)(\d+)(밀리그램|mg|정|T)?$', name)
        if m and " " not in name:
            spaced = f"{m.group(1)} {m.group(2)}{m.group(3) or ''}"
            if spaced not in name_variants:
                name_variants.append(spaced)

    return {
        "product_name": name,
        "product_code": code,
        "api_name": api,
        "name_variants": name_variants,
        "code_prefix": code[:5] if len(code) >= 5 else code,
    }

def _make_cache_key(cfg: dict) -> str:
    return f"{cfg['product_name']}|{cfg['product_code']}|{cfg['api_name']}"

# ── 유틸 ─────────────────────────────────────────────────────────────
def fmt_date(v):
    if isinstance(v, datetime): return v.strftime("%Y-%m-%d")
    s = str(v or "")
    if len(s) == 8 and s.isdigit():
        return f"{s[:4]}-{s[4:6]}-{s[6:]}"
    return s[:10] if s else ""

def fmt_date_dot(v):
    if isinstance(v, datetime): return v.strftime("%Y.%m.%d")
    s = str(v or "").strip()
    if re.match(r'\d{4}-\d{2}-\d{2}', s): return s[:10].replace('-', '.')
    if re.match(r'\d{4}\.\d{2}\.\d{2}', s): return s[:10]
    if len(s) == 8 and s.isdigit(): return f"{s[:4]}.{s[4:6]}.{s[6:]}"
    return s[:10] if s else ""

# ── BDC 배치사이즈 맵 ─────────────────────────────────────────────────
def build_batch_size_map():
    bsmap = {}
    try:
        ns = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
        with zipfile.ZipFile(get_path("bdc_100k")) as z:
            shared = []
            try:
                with z.open("xl/sharedStrings.xml") as f:
                    for si in ET.parse(f).findall(f".//{{{ns}}}si"):
                        shared.append("".join(t.text or "" for t in si.findall(f".//{{{ns}}}t")))
            except: pass
            with z.open("xl/worksheets/sheet5.xml") as f:
                tree = ET.parse(f)
            for row in tree.findall(f".//{{{ns}}}row"):
                r = int(row.attrib["r"])
                if r < 19: continue
                cells = {}
                for cell in row.findall(f"{{{ns}}}c"):
                    col = sum((ord(c)-64)*(26**i) for i,c in enumerate(reversed("".join(filter(str.isalpha, cell.attrib["r"])))))
                    t = cell.attrib.get("t","")
                    v = cell.find(f"{{{ns}}}v")
                    if v is not None and v.text:
                        cells[col] = shared[int(v.text)] if t=="s" else v.text
                val = cells.get(1, "").strip()
                if val and any(val.startswith(p) for p in ["CJ","CL","CH"]):
                    bsmap[val] = {"size": "100,000T", "flag": not bool(BATCH_NORMAL_RE.match(val))}
    except: pass

    try:
        wb = openpyxl.load_workbook(get_path("bdc_700k"), data_only=True)
        ws = wb["In-Process(API 5000154,5000183)"]
        for r in range(19, 60):
            val = str(ws.cell(r, 1).value or "").strip()
            if val and any(val.startswith(p) for p in ["CJ","CL","CH"]):
                bsmap[val] = {"size": "700,000T", "flag": not bool(BATCH_NORMAL_RE.match(val))}
    except: pass

    return bsmap

BATCH_SIZE_MAP = build_batch_size_map()
_tables_cache: dict | None = None
_tables_cache_key: str = ""

def _preload_cache():
    global _tables_cache, _tables_cache_key
    cfg = get_product_config()
    if not cfg["product_code"] and not cfg["product_name"]:
        return
    _tables_cache = build_all_tables(cfg)
    _tables_cache_key = _make_cache_key(cfg)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 표11 — 제조 배치
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def map_t11(cfg: dict):
    code_prefix = cfg["code_prefix"]

    wb_w = load_wb("wms")
    ws_w = wb_w.active
    wms_map = {}
    for r in range(4, 2360):
        if not ws_w.cell(r,1).value and not ws_w.cell(r,2).value: continue
        code = str(ws_w.cell(r,5).value or "")
        if code != code_prefix: continue
        bn   = str(ws_w.cell(r,7).value or "").strip()
        qty  = ws_w.cell(r,9).value
        unit = str(ws_w.cell(r,10).value or "")
        mfg  = fmt_date(str(ws_w.cell(r,3).value or ""))
        if bn and bn not in wms_map:
            wms_map[bn] = {"mfg_date": mfg, "batch_size": f"{int(qty):,}{unit}" if qty else "?", "wms_row": r}
    wb_w.close()

    wb_b = load_wb("batch", read_only=True)
    ws_b = wb_b.active
    rows = []
    for r, row_cells in enumerate(ws_b.iter_rows(min_row=4, max_row=624, values_only=True), start=4):
        if not row_cells[0]: continue
        code = str(row_cells[2] or "")
        if not code_prefix or code_prefix not in code: continue
        prod_batch = str(row_cells[4] or "")
        mfg_raw    = fmt_date(row_cells[6])
        mfg_year   = int(mfg_raw[:4]) if mfg_raw and mfg_raw[:4].isdigit() else 0
        if mfg_year != 2025: continue
        batch_no   = prod_batch[:-1] if prod_batch and prod_batch[-1].isalpha() else prod_batch
        market     = MARKET_MAP.get(code[:7], code)
        remark_raw = str(row_cells[9] or "")
        remark_parts = []
        if "안정성" in remark_raw or "Stability" in remark_raw: remark_parts.append("Stability")
        if "PKV" in remark_raw: remark_parts.append("PKV")
        if "sublot" in remark_raw.lower(): remark_parts.append("Sublot")
        remark = ", ".join(remark_parts) if remark_parts else "N/A"
        wms    = wms_map.get(batch_no, {})
        binfo  = BATCH_SIZE_MAP.get(batch_no)
        mfg_date   = wms.get("mfg_date", fmt_date(row_cells[6]))
        batch_size = wms.get("batch_size") or (binfo["size"] if binfo else "?")
        wms_row    = wms.get("wms_row")
        review = False; review_reason = ""
        if binfo and binfo["flag"]:
            review = True; review_reason = f"비정상 배치번호 기재: {batch_no}"
        elif not binfo:
            review = True; review_reason = "BDC 파일에서 배치번호 미확인"
        rows.append({
            "data": {"No": str(len(rows)+1), "Batch No": batch_no, "Date of Manufacture": mfg_date,
                     "Batch Size": batch_size,
                     "Packaging Unit": PACKAGING_UNIT_MAP.get(code[:7], str(row_cells[3] or "")),
                     "Product Batch No.": prod_batch, "Market": market, "Remark": remark},
            "sources": [{"file": "자사 제조배치.xlsx", "sheet": "Sheet1", "row": r},
                        *([{"file": "WMS 제조지시 리스트.xlsx", "sheet": "Sheet1", "row": wms_row}] if wms_row else [])],
            "review_required": review, "review_reason": review_reason,
        })
    wb_b.close()

    rows_by_batch: dict = {}
    for row in rows:
        bn = row["data"]["Batch No"]
        rows_by_batch.setdefault(bn, []).append(row)
    sorted_rows = []
    for bn in sorted(rows_by_batch.keys()):
        sorted_rows.extend(rows_by_batch[bn])
    for i, row in enumerate(sorted_rows):
        row["data"]["No"] = str(i + 1)
    return {"id": "t11", "title": "표11 — 제조 배치",
            "columns": ["No","Batch No","Date of Manufacture","Batch Size","Packaging Unit","Product Batch No.","Market","Remark"],
            "rows": sorted_rows}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 표16 — 변경관리
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def map_t16(cfg: dict):
    wb = load_wb("cc")
    ws = wb["Change Control List"]
    code_prefix   = cfg["code_prefix"]
    name_variants = cfg["name_variants"]

    def is_code_match(val):
        return bool(code_prefix) and code_prefix in str(val or "")

    def is_name_match(val):
        s = str(val or "")
        return any(v in s for v in name_variants)

    rows = []
    for r in range(5, ws.max_row + 1):
        if not ws.cell(r,2).value: continue
        affected = ws.cell(r,5).value
        if not is_code_match(affected) and not is_name_match(affected): continue
        cc_no      = str(ws.cell(r,2).value or "")
        title      = str(ws.cell(r,4).value or "").replace("\n"," ").replace("\t"," ")
        open_str   = fmt_date(ws.cell(r,11).value)
        complete   = ws.cell(r,12).value
        complete_s = str(complete or "").strip()
        if complete and complete_s not in ("TBD","None",""):
            close_str = fmt_date(complete)
        elif complete_s == "TBD":
            close_str = "TBD"
        else:
            close_str = "진행 중"
        review = not is_code_match(affected)
        review_reason = "제품코드 미기재 (텍스트 매칭 — 담당자 확인 필요)" if review else ""
        rows.append({
            "data": {"No": str(len(rows)+1), "Change no.": cc_no,
                     "Change Information": title, "Open/Close date": f"{open_str} / {close_str}"},
            "sources": [{"file": "변경관리 목록_2025년.xlsx", "sheet": "Change Control List", "row": r}],
            "review_required": review, "review_reason": review_reason,
        })
    return {"id": "t16", "title": "표16 — 평가기간 변경관리",
            "columns": ["No","Change no.","Change Information","Open/Close date"], "rows": rows}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 표18 — CAPA
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def map_t18(cfg: dict):
    wb = load_wb("capa")
    ws = wb["Chemical"]
    code_prefix   = cfg["code_prefix"]
    name_variants = cfg["name_variants"]
    filter_terms  = ([code_prefix] if code_prefix else []) + name_variants

    def is_match(text: str) -> bool:
        return any(t in text for t in filter_terms)

    def is_code_match(text: str) -> bool:
        return bool(code_prefix) and code_prefix in text

    rows = []
    for r in range(3, ws.max_row+1):
        capa_no = str(ws.cell(r,1).value or "").strip()
        if not capa_no or "CAPA" not in capa_no: continue
        title   = str(ws.cell(r,4).value or "").replace("\n"," ").replace("\t"," ")
        product = str(ws.cell(r,5).value or "").replace("\n"," ")
        status  = str(ws.cell(r,6).value or "").strip()
        opened  = str(ws.cell(r,7).value or "")[:12]
        planned = str(ws.cell(r,8).value or "")[:12]
        product_stripped = product.strip()
        combined = f"{capa_no} {title} {product}"
        if product_stripped != "All" and not is_match(combined): continue
        open_year = int(opened[:4]) if opened and opened[:4].isdigit() else 0
        if open_year < 2025: continue
        has_code = is_code_match(product)
        review = not has_code
        review_reason = "Related Product 필드에 제품코드 미기재" if review else ""
        if "complete" in status.lower() or "closed" in status.lower():
            close_str = f"{planned} (Complete)" if planned else "Complete"
        elif planned:
            close_str = planned
        else:
            close_str = "진행 중"
            review = True
            review_reason = (review_reason + "; " if review_reason else "") + "종료일 미기재 — 확인 필요"
        rows.append({
            "data": {"No": str(len(rows)+1), "CAPA No.": capa_no,
                     "Summary": title, "Open/Close date": f"{opened} / {close_str}"},
            "sources": [{"file": "List of CAPA (2025)_QA1.xlsx", "sheet": "Chemical", "row": r}],
            "review_required": review, "review_reason": review_reason,
        })
    return {"id": "t18", "title": "표18 — 평가기간 CAPA",
            "columns": ["No","CAPA No.","Summary","Open/Close date"], "rows": rows}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 표19 — 불만
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def map_t19(cfg: dict):
    wb = load_wb("complaint")
    ws = wb.active
    code_prefix   = cfg["code_prefix"]
    name_variants = cfg["name_variants"]
    filter_terms  = name_variants + ([code_prefix] if code_prefix else [])

    rows = []
    for r in range(2, ws.max_row+1):
        product = str(ws.cell(r,5).value or "")
        if not any(t in product for t in filter_terms): continue
        recv         = fmt_date_dot(ws.cell(r,1).value)
        comp_no      = str(ws.cell(r,4).value or "")
        comp_type_kr = str(ws.cell(r,11).value or "").replace("\n"," ").strip()
        invest       = str(ws.cell(r,15).value or "").replace("\n"," ").strip()
        capa_act     = str(ws.cell(r,16).value or "").replace("\n"," ").strip()
        capa_nos     = str(ws.cell(r,18).value or "").replace("\n",", ").strip()
        close        = fmt_date_dot(ws.cell(r,14).value)
        comp_type_en = COMPLAINT_TYPE_EN.get(comp_type_kr, "")
        contents_line = f"{comp_type_en} {comp_type_kr}" if comp_type_en else comp_type_kr
        action_text   = capa_act or invest
        summary_parts = [f"Contents내용", contents_line]
        if action_text:
            summary_parts += ["Corrective Action조치내용", action_text]
        if capa_nos and capa_nos not in ("N/A",""):
            summary_parts.append(f"[CAPA: {capa_nos}]")
        summary  = "\n".join(summary_parts)
        date_str = f"{recv}/\n{close}" if close else f"{recv}"
        rows.append({
            "data": {"No": str(len(rows)+1), "Complaint No.": comp_no,
                     "Contents / Corrective Action": summary,
                     "Receive / Complete Date": date_str},
            "sources": [{"file": "소비자불만 리스트.xlsx", "sheet": "Sheet1", "row": r}],
            "review_required": False, "review_reason": "",
        })
    return {"id": "t19", "title": "표19 — 불만 (Complaints)",
            "columns": ["No","Complaint No.","Contents / Corrective Action","Receive / Complete Date"],
            "rows": rows}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 표25 — SCAR
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def map_scar(cfg: dict):
    wb = load_wb("scar")
    ws = wb["2025"]
    name_variants = cfg["name_variants"]
    code_prefix   = cfg["code_prefix"]
    filter_terms  = name_variants + ([code_prefix] if code_prefix else [])

    rows = []
    for r in range(2, ws.max_row + 1):
        scar_no = str(ws.cell(r, 1).value or "").strip()
        if not scar_no: continue
        supplier    = str(ws.cell(r, 2).value or "").replace("\n", " ").strip()
        material    = str(ws.cell(r, 3).value or "").strip()
        obs         = str(ws.cell(r, 5).value or "").strip()
        send_dt     = ws.cell(r, 6).value
        complete_dt = ws.cell(r, 8).value
        capa_act    = str(ws.cell(r, 9).value or "").strip()
        impact      = str(ws.cell(r, 10).value or "").strip()
        status      = str(ws.cell(r, 11).value or "").strip()

        combined = material + obs + capa_act
        if filter_terms and not any(t in combined for t in filter_terms): continue

        send_str   = fmt_date(send_dt)
        complete_s = str(complete_dt or "").strip()
        if complete_dt and complete_s not in ("TBD", "None", ""):
            close_str = fmt_date(complete_dt)
        elif complete_s == "TBD":
            close_str = "TBD"
        elif "완료" in status or "complete" in status.lower():
            close_str = "완료"
        else:
            close_str = "진행 중"

        content_parts = ["1. 발견 내용", obs]
        if capa_act:
            content_parts += ["", "2. 조치 내용", capa_act]
        if impact:
            content_parts += ["", "3. 평가", impact]
        contents = "\n".join(content_parts)

        has_direct = any(t in material for t in filter_terms) if filter_terms else False
        review = not has_direct
        review_reason = "원자재명에 제품 미기재 — 담당자 확인 필요" if review else ""

        rows.append({
            "data": {
                "No": scar_no,
                "Company": supplier,
                "Contents/ Corrective Action/ Evaluation": contents,
                "Send date/\nComplete Date": f"{send_str} /\n{close_str}",
            },
            "sources": [{"file": "SCAR DB_2025.xlsx", "sheet": "2025", "row": r}],
            "review_required": review,
            "review_reason": review_reason,
        })
    wb.close()
    return {"id": "t_scar", "title": "표25 — SCAR (공급업체 시정조치)",
            "columns": ["No", "Company", "Contents/ Corrective Action/ Evaluation", "Send date/\nComplete Date"],
            "rows": rows}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 일탈 (Deviation Log)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def map_deviation(cfg: dict):
    wb = load_wb("deviation")
    ws = wb["2025"]
    name_variants = cfg["name_variants"]
    code_prefix   = cfg["code_prefix"]

    rows = []
    for r in range(4, ws.max_row + 1):          # 헤더 row3, 데이터 row4~
        dev_no = str(ws.cell(r, 1).value or "").strip()
        if not dev_no: continue

        level    = str(ws.cell(r, 2).value or "").strip()
        title    = str(ws.cell(r, 3).value or "").strip().split("\n")[0]  # 영문 제목만
        open_dt  = ws.cell(r, 8).value
        close_dt = ws.cell(r, 10).value
        status   = str(ws.cell(r, 11).value or "").strip()
        prod_nm  = str(ws.cell(r, 15).value or "").strip()   # 제품명
        batch_mk = str(ws.cell(r, 16).value or "").strip()   # 제조번호(마켓)

        # 필터: 제품명에 name_variants 포함, 없으면 code_prefix로 시도
        if name_variants:
            if not any(v in prod_nm for v in name_variants): continue
        elif code_prefix:
            if code_prefix not in prod_nm and code_prefix not in batch_mk: continue

        open_str  = fmt_date(open_dt)
        close_str = fmt_date(close_dt) if close_dt else ("진행 중" if status != "종료" else "")

        content = f"[{level}] {title}"
        date_str = f"{open_str} /\n{close_str}" if close_str else open_str

        review = (status != "종료")
        review_reason = "미종료 일탈 — 진행 상태 확인 필요" if review else ""

        rows.append({
            "data": {
                "No": str(len(rows) + 1),
                "일탈 번호": dev_no,
                "일탈 내용 / 조치사항/ 평가": content,
                "개시일/완료일": date_str,
            },
            "sources": [{"file": "Deviation Log 2025.xlsx", "sheet": "2025", "row": r}],
            "review_required": review,
            "review_reason": review_reason,
        })
    wb.close()
    return {"id": "t_deviation", "title": "일탈 (Deviation Log 2025)",
            "columns": ["No", "일탈 번호", "일탈 내용 / 조치사항/ 평가", "개시일/완료일"],
            "rows": rows}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 반품 내역
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def map_banpum(cfg: dict):
    import pyxlsb
    code_prefix   = cfg["code_prefix"]
    name_variants = cfg["name_variants"]
    path = get_path("banpum")

    groups: dict = {}   # 반품접수번호 → {date, unit, qty, reason, batches}
    with pyxlsb.open_workbook(str(path)) as wb:
        with wb.get_sheet("Sheet1") as ws:
            for ri, row in enumerate(ws.rows(), start=1):
                if ri <= 4: continue   # 제목·빈행·헤더·공백 스킵
                vals = [c.v for c in row]
                if len(vals) < 16: continue
                ref_no   = str(vals[1] or "").strip()
                if not ref_no: continue
                prod_code = str(vals[3] or "").strip()
                prod_nm   = str(vals[4] or "").strip()
                unit      = str(vals[5] or "").strip()
                batch_no  = str(vals[6] or "").strip()
                qty       = vals[8] or 0
                ret_date  = str(vals[13] or "").strip()
                proc_date = str(vals[14] or "").strip()
                reason    = str(vals[15] or "").strip()

                # 필터: 제품코드 우선, 없으면 제품명
                if code_prefix:
                    if code_prefix not in prod_code: continue
                elif name_variants:
                    if not any(v in prod_nm for v in name_variants): continue

                if ref_no not in groups:
                    groups[ref_no] = {"date": ret_date, "proc": proc_date,
                                      "unit": unit, "reason": reason,
                                      "qty": 0, "batches": [], "row": ri}
                groups[ref_no]["qty"] += float(qty) if qty else 0
                if batch_no:
                    groups[ref_no]["batches"].append(batch_no)

    rows = []
    for ref_no, g in groups.items():
        qty_str = f"{int(g['qty'])} {g['unit']}" if g["qty"] == int(g["qty"]) else f"{g['qty']} {g['unit']}"
        batch_str = ", ".join(sorted(set(g["batches"])))
        review = not g["proc"]
        review_reason = "처리일자 없음 — 처분 결과 확인 필요" if review else ""
        rows.append({
            "data": {
                "No": str(len(rows) + 1),
                "반품일자": g["date"],
                "반품수량": qty_str,
                "반품사유": g["reason"],
                "처분결과": g["proc"] or "",
                "제조번호": batch_str,
            },
            "sources": [{"file": "반품 내역.xlsb", "sheet": "Sheet1", "row": g["row"]}],
            "review_required": review,
            "review_reason": review_reason,
        })
    return {"id": "t_banpum", "title": "반품 내역 (2025)",
            "columns": ["No", "반품일자", "반품수량", "반품사유", "처분결과", "제조번호"],
            "rows": rows}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 표22 — 원자재(API) 수입 시험 현황
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def map_t22_api(cfg: dict):
    import pyxlsb
    api_name = cfg.get("api_name", "").strip()
    path = get_path("rmt")
    with pyxlsb.open_workbook(str(path)) as wb:
        with wb.get_sheet("Sheet1") as ws:
            all_rows = list(ws.rows())

    def api_matches(name: str) -> bool:
        if not api_name:
            return True  # no filter → show all APIs
        n = name.strip().lower()
        return n == api_name.lower() or api_name.lower() in n

    groups: dict = {}
    for ri, row in enumerate(all_rows[3:], start=4):
        vals = [c.v for c in row]
        if len(vals) < 5: continue
        name = str(vals[3] or "").strip()
        if not api_matches(name): continue
        dt = str(vals[0] or "")
        if "2025" not in dt: continue
        code   = str(vals[2] or "")
        status = str(vals[4] or "").strip()
        key = (code, name)
        if key not in groups:
            groups[key] = {"total": 0, "approved": 0, "rejected": 0, "other": 0, "first_row": ri}
        g = groups[key]
        g["total"] += 1
        if status == "적합":
            g["approved"] += 1
        elif "부적합" in status:
            g["rejected"] += 1
        else:
            g["other"] += 1

    rows = []
    for (code, name), g in groups.items():
        review = g["rejected"] > 0
        review_reason = f"부적합 {g['rejected']}건 발생 — 확인 필요" if review else ""
        rows.append({
            "data": {
                "No": str(len(rows) + 1),
                "Name of API": name,
                "Code No.": code,
                "Total Receipt": str(g["total"]),
                "Approved": str(g["approved"]),
                "Under Test": str(g["other"]),
                "Reject": str(g["rejected"]),
            },
            "sources": [{"file": "원자재 시험성적.xlsb", "sheet": "Sheet1", "row": g["first_row"]}],
            "review_required": review,
            "review_reason": review_reason,
        })
    return {"id": "t22_api", "title": "표22 — API 수입 시험 현황 (2025)",
            "columns": ["No", "Name of API", "Code No.", "Total Receipt", "Approved", "Under Test", "Reject"],
            "rows": rows}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 완제품 QC 출하 현황
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def map_fpt(cfg: dict):
    import pyxlsb
    code_prefix = cfg["code_prefix"]
    path = get_path("fpt")
    with pyxlsb.open_workbook(str(path)) as wb:
        with wb.get_sheet("Sheet1") as ws:
            all_rows = list(ws.rows())

    batch_entries: dict = {}
    for ri, row in enumerate(all_rows[3:], start=4):
        vals = [c.v for c in row]
        if len(vals) < 15: continue
        code = str(vals[9] or "")
        if code_prefix and code_prefix not in code: continue
        dt = str(vals[8] or "")
        if "2025" not in dt: continue
        batch = str(vals[14] or "")
        if not batch: continue
        status = str(vals[2] or "")
        market = MARKET_MAP.get(code, code)  # fallback to raw code
        batch_entries.setdefault(batch, []).append((ri, market, status, dt[:10]))

    rows = []
    for batch_no, entries in sorted(batch_entries.items()):
        markets  = sorted(set(e[1] for e in entries))
        statuses = sorted(set(e[2] for e in entries))
        dates    = sorted(set(e[3] for e in entries))
        all_approved = all(e[2] == "출하승인" for e in entries)
        review = not all_approved
        review_reason = "출하 미승인 배치 포함 — 확인 필요" if review else ""
        rows.append({
            "data": {
                "No": str(len(rows) + 1),
                "Batch No": batch_no,
                "Market": ", ".join(markets),
                "시험상태": ", ".join(statuses),
                "QC 승인일": ", ".join(dates),
            },
            "sources": [{"file": "완제품 시험성적.xlsb", "sheet": "Sheet1", "row": entries[0][0]}],
            "review_required": review,
            "review_reason": review_reason,
        })
    return {"id": "t_fpt", "title": "완제품 QC 출하 현황 (2025)",
            "columns": ["No", "Batch No", "Market", "시험상태", "QC 승인일"],
            "rows": rows}

def build_all_tables(cfg: dict) -> dict:
    return {
        "t11":        map_t11(cfg),
        "t16":        map_t16(cfg),
        "t18":        map_t18(cfg),
        "t19":        map_t19(cfg),
        "t_scar":     map_scar(cfg),
        "t_deviation": map_deviation(cfg),
        "t_banpum":   map_banpum(cfg),
    }

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 소스 헤더 자동 감지
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _find_header_row(ws, max_search: int = 8, max_col: int = 18) -> int:
    from datetime import datetime as _dt
    best_row, best_count = 1, 0
    for r in range(1, max_search + 1):
        vals = [ws.cell(r, c).value for c in range(1, max_col + 1)]
        if any(isinstance(v, (int, float)) or isinstance(v, _dt) for v in vals):
            continue
        ne = sum(1 for v in vals if v)
        if ne > best_count:
            best_count = ne; best_row = r
    return best_row

def _find_header_row_iterrows(first_rows: list[tuple]) -> int:
    from datetime import datetime as _dt
    best_idx, best_count = 0, 0
    for idx, row_vals in enumerate(first_rows):
        if any(isinstance(v, (int, float)) or isinstance(v, _dt) for v in row_vals):
            continue
        ne = sum(1 for v in row_vals if v)
        if ne > best_count:
            best_count = ne; best_idx = idx
    return best_idx

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Word 양식 채우기
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
WORD_TABLE_MAP = {
    10: ("t11",          ["No","Batch No","Date of Manufacture","Batch Size","Packaging Unit","Product Batch No.","Market","Remark"]),
    15: ("t16",          ["No","Change no.","Change Information","Open/Close date"]),
    17: ("t18",          ["No","CAPA No.","Summary","Open/Close date"]),
    18: ("t19",          ["No","Complaint No.","Contents / Corrective Action","Receive / Complete Date"]),
    25: ("t_scar",       ["No","Company","Contents/ Corrective Action/ Evaluation","Send date/\nComplete Date"]),
     7: ("t_deviation",  ["No","일탈 번호","일탈 내용 / 조치사항/ 평가","개시일/완료일"]),
    12: ("t_banpum",     ["No","반품일자","반품수량","반품사유","처분결과"]),
}

def _fill_table(word_tbl, data_rows: list[dict], columns: list[str]):
    from docx.oxml.ns import qn
    from lxml import etree
    XML_SPACE_NS = "http://www.w3.org/XML/1998/namespace"

    def _write_tc(tc, text: str):
        paras = tc.findall(qn("w:p"))
        if not paras:
            return
        p = paras[0]
        lines = text.split("\n") if "\n" in text else [text]
        existing_r = p.find(qn("w:r"))
        rPr = copy.deepcopy(existing_r.find(qn("w:rPr"))) if existing_r is not None else None
        for r in p.findall(qn("w:r")):
            p.remove(r)
        if not text:
            return
        r_elem = etree.SubElement(p, qn("w:r"))
        if rPr is not None:
            r_elem.insert(0, rPr)
        for i, line in enumerate(lines):
            if i > 0:
                etree.SubElement(r_elem, qn("w:br"))
            t = etree.SubElement(r_elem, qn("w:t"))
            t.text = line
            if line != line.strip():
                t.set(f"{{{XML_SPACE_NS}}}space", "preserve")

    def add_row_copy():
        last_tr = word_tbl.rows[-1]._tr
        new_tr  = copy.deepcopy(last_tr)
        for tc in new_tr.findall(qn("w:tc")):
            for t in tc.findall(f".//{qn('w:t')}"):
                t.text = ""
            for vm in tc.findall(f".//{qn('w:vMerge')}"):
                parent = vm.getparent()
                if parent is not None:
                    parent.remove(vm)
        last_tr.addnext(new_tr)
        return word_tbl.rows[-1]

    existing = list(word_tbl.rows[1:])
    for i, row_data in enumerate(data_rows):
        tbl_row = existing[i] if i < len(existing) else add_row_copy()
        raw_tcs = tbl_row._tr.findall(qn("w:tc"))
        for col_idx, col_name in enumerate(columns):
            if col_idx >= len(raw_tcs):
                break
            tc = raw_tcs[col_idx]
            vm = tc.find(f".//{qn('w:vMerge')}")
            if vm is not None:
                vm_val = vm.get(qn("w:val"), "continue")
                if vm_val != "restart":
                    continue
            val = str(row_data.get("data", {}).get(col_name, "") or "")
            _write_tc(tc, val)

def generate_filled_word(tables: dict) -> Path:
    from docx import Document
    import tempfile
    tpl_path = get_path("template")
    doc = Document(str(tpl_path))
    for tbl_idx, (tid, columns) in WORD_TABLE_MAP.items():
        if tbl_idx >= len(doc.tables): continue
        td = tables.get(tid)
        if not td: continue
        _fill_table(doc.tables[tbl_idx], td["rows"], columns)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=str(UPLOADS))
    doc.save(tmp.name)
    tmp.close()
    return Path(tmp.name)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# API 엔드포인트
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
@app.get("/api/tables")
def get_tables(product_name: str = "", product_code: str = "", api_name: str = "", refresh: bool = False):
    global _tables_cache, _tables_cache_key
    cfg = get_product_config(product_name, product_code, api_name)
    if not cfg["product_code"]:
        return {"error": "no_product", "message": "제품코드를 먼저 입력하세요."}
    cache_key = _make_cache_key(cfg)
    if _tables_cache is None or refresh or cache_key != _tables_cache_key:
        _tables_cache = build_all_tables(cfg)
        _tables_cache_key = cache_key
    return _tables_cache


@app.post("/api/upload")
async def upload_files(
    product_name: str      = Form(""),
    product_code: str      = Form(""),
    api_name:     str      = Form(""),
    batch:        UploadFile = File(None),
    wms:          UploadFile = File(None),
    cc:           UploadFile = File(None),
    capa:         UploadFile = File(None),
    complaint:    UploadFile = File(None),
    scar:         UploadFile = File(None),
    rmt:          UploadFile = File(None),
    fpt:          UploadFile = File(None),
    template:     UploadFile = File(None),
):
    if product_name or product_code:
        info = {"product_name": product_name, "product_code": product_code, "api_name": api_name}
        PRODUCT_INFO_FILE.write_text(json.dumps(info, ensure_ascii=False, indent=2), encoding="utf-8")

    uploaded = []
    pairs = [("batch", batch), ("wms", wms), ("cc", cc),
             ("capa", capa), ("complaint", complaint), ("scar", scar),
             ("rmt", rmt), ("fpt", fpt), ("template", template)]
    for key, uf in pairs:
        if uf and uf.filename:
            data = await uf.read()
            dest = UPLOAD_FILES[key]
            dest.write_bytes(data)
            uploaded.append(key)

    global _tables_cache
    _tables_cache = None
    return {"uploaded": uploaded, "count": len(uploaded)}


@app.post("/api/upload/clear")
def clear_uploads():
    global _tables_cache
    for key, path in UPLOAD_FILES.items():
        if path.exists():
            path.unlink()
    if PRODUCT_INFO_FILE.exists():
        PRODUCT_INFO_FILE.unlink()
    _tables_cache = None
    return {"cleared": True}


@app.get("/api/upload/status")
def upload_status():
    return {
        key: {
            "uploaded": UPLOAD_FILES[key].exists(),
            "filename": UPLOAD_FILES[key].name if UPLOAD_FILES[key].exists() else None,
        }
        for key in ["batch", "wms", "cc", "capa", "complaint", "scar", "rmt", "fpt", "template"]
    }


@app.get("/api/download-word")
def download_word():
    global _tables_cache, _tables_cache_key
    cfg = get_product_config()
    if not cfg["product_name"] and not cfg["product_code"]:
        return {"error": "제품 정보가 설정되지 않았습니다."}
    cache_key = _make_cache_key(cfg)
    if _tables_cache is None or cache_key != _tables_cache_key:
        _tables_cache = build_all_tables(cfg)
        _tables_cache_key = cache_key
    try:
        out_path = generate_filled_word(_tables_cache)
        return FileResponse(
            path=str(out_path),
            filename="PPQR_완성본.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        return {"error": str(e)}


@app.get("/api/sample/{key}")
def download_sample(key: str):
    path = SAMPLE_FILES.get(key)
    if not path or not path.exists():
        return {"error": "파일 없음"}
    return FileResponse(path=str(path), filename=path.name)


@app.get("/api/source")
def get_source(file: str = Query(...), sheet: str = Query(...), row: int = Query(...)):
    FILE_KEY_MAP = {
        "자사 제조배치.xlsx":            "batch",
        "WMS 제조지시 리스트.xlsx":       "wms",
        "변경관리 목록_2025년.xlsx":      "cc",
        "List of CAPA (2025)_QA1.xlsx": "capa",
        "소비자불만 리스트.xlsx":          "complaint",
        "SCAR DB_2025.xlsx":            "scar",
        "원자재 시험성적.xlsb":           "rmt",
        "완제품 시험성적.xlsb":           "fpt",
    }
    key = FILE_KEY_MAP.get(file)
    if not key:
        return {"error": f"Unknown file: {file}"}

    try:
        start  = max(1, row - 7)
        end    = row + 7
        SEARCH = 8

        if key == "batch":
            wb = load_wb("batch", read_only=True)
            ws = wb.active
            first_rows: list[tuple] = []
            rows_data = []
            for r_num, row_vals in enumerate(ws.iter_rows(min_row=1, max_row=end, values_only=True), start=1):
                if r_num <= SEARCH:
                    first_rows.append(row_vals[:18])
                if r_num < start:
                    continue
                cells = [str(v or "") for v in row_vals[:18]]
                rows_data.append({"row_num": r_num, "cells": cells, "is_match": r_num == row})
            wb.close()
            hdr_idx = _find_header_row_iterrows(first_rows)
            raw     = list(first_rows[hdr_idx]) if first_rows else []
            headers = [str(v or "") for v in raw] if any(raw) else [f"열{i+1}" for i in range(18)]
        elif key in ("capa", "scar"):
            wb  = load_wb(key)
            ws  = wb[sheet]
            max_col = min(ws.max_column, 18)
            end = min(ws.max_row, end)
            hdr = _find_header_row(ws, max_search=SEARCH, max_col=max_col)
            headers   = [str(ws.cell(hdr, c).value or "") for c in range(1, max_col+1)]
            rows_data = [{"row_num": r, "cells": [str(ws.cell(r, c).value or "") for c in range(1, max_col+1)], "is_match": r == row} for r in range(start, end+1)]
        elif key in ("rmt", "fpt"):
            import pyxlsb
            path = get_path(key)
            with pyxlsb.open_workbook(str(path)) as xlwb:
                with xlwb.get_sheet("Sheet1") as xlws:
                    all_xlsb = list(xlws.rows())
            max_col = 18
            hdr_row = all_xlsb[2] if len(all_xlsb) > 2 else []
            headers   = [str(c.v or "") for c in hdr_row[:max_col]]
            rows_data = []
            for r_num, xlrow in enumerate(all_xlsb[start - 1:end], start=start):
                cells = [str(c.v or "") for c in xlrow[:max_col]]
                rows_data.append({"row_num": r_num, "cells": cells, "is_match": r_num == row})
        else:
            wb  = load_wb(key)
            ws  = wb.active
            max_col = min(ws.max_column, 18)
            end = min(ws.max_row, end)
            hdr = _find_header_row(ws, max_search=SEARCH, max_col=max_col)
            headers   = [str(ws.cell(hdr, c).value or "") for c in range(1, max_col+1)]
            rows_data = [{"row_num": r, "cells": [str(ws.cell(r, c).value or "") for c in range(1, max_col+1)], "is_match": r == row} for r in range(start, end+1)]

        return {"headers": headers, "rows": rows_data, "match_row": row, "file": file, "sheet": sheet}
    except Exception as e:
        return {"error": str(e)}


@app.get("/api/health")
def health():
    return {
        "status": "ok",
        "sample_files":   {k: v.exists() for k, v in SAMPLE_FILES.items()},
        "uploaded_files":  {k: v.exists() for k, v in UPLOAD_FILES.items()},
        "product_info": json.loads(PRODUCT_INFO_FILE.read_text(encoding="utf-8"))
                        if PRODUCT_INFO_FILE.exists() else None,
    }
